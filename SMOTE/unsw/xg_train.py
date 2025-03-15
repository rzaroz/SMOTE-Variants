import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
import matplotlib.pyplot as plt
from xgboost import XGBClassifier
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, OneHotEncoder, Normalizer
from sklearn.metrics import classification_report, confusion_matrix, accuracy_score, ConfusionMatrixDisplay

# Load dataset
title = "UNSW-NB15 - XGBOOST - SMOTE"
df_src = "report/SMOTE-UNSW_NB15.csv"
df = pd.read_csv(df_src)

# Preprocessing
X = df.drop(["attack_cat"], axis=1)

# Normalize data
scaler = Normalizer()
X = scaler.fit_transform(X)

label_encoder = LabelEncoder()
y = label_encoder.fit_transform(df["attack_cat"])

X_train, X_test, y_train, y_test = train_test_split(X, y, random_state=42, test_size=0.2)

print("Start training!")
model = XGBClassifier()
model.fit(X_train, y_train)
print("Done!")
y_hat = model.predict(X_test)
print(y_hat)
clf_report = classification_report(y_test, y_hat, target_names=label_encoder.classes_, output_dict=True)
print("CLF report:")
print(clf_report)
print("-----------------------------")
acc = accuracy_score(y_test, y_hat)
print("Accuracy: ", acc)
print("-----------------------------")
conf_mat = confusion_matrix(y_test, y_hat)
disp = ConfusionMatrixDisplay(confusion_matrix=conf_mat, display_labels=label_encoder.classes_)
disp.plot(cmap=plt.cm.Blues)
plt.title(title)
plt.xticks(rotation=40, fontsize=5)
plt.yticks(fontsize=5)
plt.show()


# Create table of clf report
clf_table = pd.DataFrame(clf_report).transpose()
doc = Document()
doc.add_heading(title, level=1)

table = doc.add_table(rows=clf_table.shape[0]+1, cols=clf_table.shape[1]+1)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Class/Metric'
for i, col_name in enumerate(clf_table.columns):
    hdr_cells[i+1].text = col_name

for i, (index, row) in enumerate(clf_table.iterrows()):
    row_cells = table.rows[i+1].cells
    row_cells[0].text = str(index)
    for j, value in enumerate(row):
        row_cells[j+1].text = f"{value:.2f}" if isinstance(value, (float, np.float64)) else str(value)

for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

for row in table.rows:
    for cell in row.cells:
        cell_xml = cell._tc.get_or_add_tcPr()
        borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:top w:val="single" w:sz="6"/>'
                            r'<w:left w:val="single" w:sz="6"/>'
                            r'<w:bottom w:val="single" w:sz="6"/>'
                            r'<w:right w:val="single" w:sz="6"/>'
                            r'</w:tcBorders>')
        cell_xml.append(borders)

doc.add_paragraph("------------------------------------------------------------------")
doc.add_heading("Accuracy", level=1)
doc.add_paragraph(acc.__str__())

doc.save(f'report/{title}.docx')
print("File created!")
